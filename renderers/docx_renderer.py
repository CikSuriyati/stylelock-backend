"""
DOCX renderer.

Builds a Word document programmatically from a StructuredDocument + the
GADING ruleset, applying Word *styles* (not CSS). This is what makes the
download match the web view: we don't try to convert HTML to Word — we
re-render from the same source using python-docx primitives.

If a journal template (.dotx / .docx) is supplied, we open that as the
base so all of its built-in styles, headers, and section breaks are
preserved. Otherwise we synthesize the styles from the ruleset.
"""

from __future__ import annotations

import os
import re
from typing import Dict, Any, Optional, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .schema import (
    StructuredDocument,
    HeadingBlock,
    ParagraphBlock,
    TableBlock,
    FigureBlock,
    ReferenceBlock,
    EquationBlock,
    RunSpan,
    normalize_input,
)


# ---------- alignment helpers ----------

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _line_spacing_value(spec: str) -> float:
    return {"single": 1.0, "1.5": 1.5, "double": 2.0}.get(spec, 1.0)


def _apply_runs(paragraph, runs: List[RunSpan], default_font: str, default_size: int):
    for r in runs:
        run = paragraph.add_run(r.text)
        run.font.name = default_font
        run.font.size = Pt(default_size)
        run.bold = r.bold
        run.italic = r.italic
        run.underline = r.underline
        # python-docx doesn't expose superscript/subscript directly on Font
        if r.superscript:
            run.font.superscript = True
        if r.subscript:
            run.font.subscript = True


def _add_paragraph_text(paragraph, text: str | None, runs: List[RunSpan], font: str, size: int):
    if runs:
        _apply_runs(paragraph, runs, font, size)
    elif text:
        run = paragraph.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)


# ---------- table border helpers (XML) ----------

def _set_cell_border(cell, *, top: bool = False, bottom: bool = False, left: bool = False, right: bool = False):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for side, want in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        existing = tcBorders.find(qn(f"w:{side}"))
        if existing is not None:
            tcBorders.remove(existing)
        if want:
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "6")     # 0.75pt
            b.set(qn("w:color"), "000000")
            tcBorders.append(b)


# ---------- main renderer ----------

class DocxRenderer:
    def __init__(self, ruleset: Dict[str, Any], template_path: Optional[str] = None):
        self.ruleset = ruleset or {}
        fonts = self.ruleset.get("fonts", {}) or {}
        self.main_font = (fonts.get("main_text") or {}).get("family", "Times New Roman")
        self.main_size = (fonts.get("main_text") or {}).get("size_pt", 10)
        self.table_font = (fonts.get("table_text") or {}).get("family", self.main_font)
        self.table_size = (fonts.get("table_text") or {}).get("size_pt", 8)

        spacing = self.ruleset.get("spacing", {}) or {}
        self.line_spacing = _line_spacing_value(spacing.get("main_text_line_spacing", "single"))
        self.para_gap = spacing.get("paragraph_spacing_before_after_pt", 0)

        self.headings = self.ruleset.get("headings", {}) or {}
        self.tables_cfg = self.ruleset.get("tables", {}) or {}
        self.refs_cfg = self.ruleset.get("references", {}) or {}

        # Base document — template if supplied, else blank
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
            # Preserve GADING branding elements from the template:
            #   Indices 0-5: Header / Title / Author / Affiliation placeholders
            #   Index 7:     GJSS Logo / Article Info layout table
            # Remove everything else so we can render fresh body content after them.
            body = self.doc.element.body
            children = list(body.iterchildren())
            keep_indices = {0, 1, 2, 3, 4, 5, 7}  # branding slots + logo table

            # The template has two sections:
            #   - An inner sectPr (inside a paragraph around element 92) that carries all
            #     headerReference / footerReference entries including the first-page logo header.
            #   - The final sectPr at the end of the body which has <w:titlePg/> but
            #     NO headerReference entries (it inherits from the inner sectPr).
            # When we strip the inner sectPr paragraph the only remaining sectPr has no
            # header references, so the logo disappears.  Fix: copy the references into the
            # final sectPr BEFORE removing anything.
            from copy import deepcopy
            final_sectPr = body.find(qn("w:sectPr"))
            if final_sectPr is not None:
                existing_types = {
                    el.get(qn("w:type"))
                    for el in final_sectPr.findall(qn("w:headerReference"))
                }
                for child_el in children:
                    if child_el.tag == qn("w:sectPr"):
                        continue
                    for inner in child_el.findall(".//" + qn("w:sectPr")):
                        for ref in inner.findall(qn("w:headerReference")):
                            ref_type = ref.get(qn("w:type"))
                            if ref_type not in existing_types:
                                final_sectPr.insert(0, deepcopy(ref))
                                existing_types.add(ref_type)
                        for ref in inner.findall(qn("w:footerReference")):
                            final_sectPr.insert(0, deepcopy(ref))

            for idx, child in enumerate(children):
                # Always keep sectPr (page layout) at the very end
                if child.tag == qn("w:sectPr"):
                    continue
                if idx not in keep_indices:
                    body.remove(child)

            # Store the template placeholder paragraphs so build() can fill them
            # instead of appending new paragraphs below the template structure.
            # After removal, self.doc.paragraphs order matches the kept body elements:
            #   [0] empty  [1] empty
            #   [2] title placeholder  [3] author  [4] affil1  [5] affil2
            self._has_template = True
        else:
            self.doc = Document()
            self._has_template = False
            self._set_default_style()

    # ----- style setup (when no template) -----

    def _set_default_style(self):
        style = self.doc.styles["Normal"]
        style.font.name = self.main_font
        style.font.size = Pt(self.main_size)
        pf = style.paragraph_format
        pf.line_spacing = self.line_spacing
        pf.space_before = Pt(self.para_gap)
        pf.space_after = Pt(self.para_gap)

    # ----- template placeholder helpers -----

    @staticmethod
    def _fill_para(para, text: str,
                   font_name: str = "", font_size_pt: Optional[float] = None,
                   bold: bool = False, italic: bool = False,
                   alignment: Optional[str] = None):
        """Replace all runs in a template placeholder with new text.

        Applies explicit character formatting so the result matches the
        ruleset regardless of what the template's paragraph style says.
        """
        p_elem = para._element
        for r in list(p_elem.findall(qn("w:r"))):
            p_elem.remove(r)
        for hl in list(p_elem.findall(qn("w:hyperlink"))):
            p_elem.remove(hl)
        if text:
            run = para.add_run(text)
            if font_name:
                run.font.name = font_name
            if font_size_pt:
                run.font.size = Pt(font_size_pt)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        if alignment and alignment in _ALIGN_MAP:
            para.alignment = _ALIGN_MAP[alignment]

    @staticmethod
    def _fill_cell_text(cell, paragraphs_text: List[str]):
        """Replace ALL paragraph content in a table cell with new lines of text.
        Reuses the existing paragraph elements to keep cell formatting/borders.
        """
        existing = cell.paragraphs
        # Clear every paragraph run
        for p in existing:
            for r in list(p._element.findall(qn("w:r"))):
                p._element.remove(r)
            for hl in list(p._element.findall(qn("w:hyperlink"))):
                p._element.remove(hl)

        for i, txt in enumerate(paragraphs_text):
            if i < len(existing):
                if txt:
                    existing[i].add_run(txt)
            else:
                # Need an extra paragraph
                from docx.oxml import OxmlElement
                new_p = OxmlElement("w:p")
                cell._tc.append(new_p)
                from docx.text.paragraph import Paragraph as _Para
                p = _Para(new_p, cell._tc)
                if txt:
                    p.add_run(txt)

    def _inject_template_front_matter(self, md):
        """Fill the template's built-in placeholder paragraphs and the ARTICLE
        INFO table with manuscript front-matter.  Called instead of
        _render_front_matter when a GADING/MJCET template is loaded."""
        tb = self.ruleset.get("title_block", {}) or {}
        title_sz  = (tb.get("title",       {}) or {}).get("font_size_pt", 17)
        title_bold= (tb.get("title",       {}) or {}).get("bold", False)
        author_sz = (tb.get("author",      {}) or {}).get("font_size_pt", 13)
        affil_sz  = (tb.get("affiliation", {}) or {}).get("font_size_pt",  8)

        paras = self.doc.paragraphs
        # para[2] = title placeholder
        if len(paras) > 2:
            self._fill_para(paras[2], md.title or "",
                            font_name=self.main_font, font_size_pt=title_sz,
                            bold=title_bold, alignment="center")
        # para[3] = author line (all authors joined)
        if len(paras) > 3:
            self._fill_para(paras[3], ",  ".join(md.authors) if md.authors else "",
                            font_name=self.main_font, font_size_pt=author_sz,
                            alignment="center")
        # para[4] = affiliation 1, para[5] = affiliation 2
        for slot, aff in enumerate(md.affiliations[:2]):
            if len(paras) > 4 + slot:
                self._fill_para(paras[4 + slot], aff,
                                font_name=self.main_font, font_size_pt=affil_sz,
                                italic=True, alignment="center")
        # Clear any unused affiliation slots
        for slot in range(len(md.affiliations), 2):
            if len(paras) > 4 + slot:
                self._fill_para(paras[4 + slot], "",
                                font_name=self.main_font, font_size_pt=affil_sz)

        if not self.doc.tables:
            return
        tbl = self.doc.tables[0]

        # Abstract → row 1, col 2 (vertically merged with row 2, col 2)
        if len(tbl.rows) > 1 and len(tbl.rows[1].cells) > 2:
            abstract_cell = tbl.rows[1].cells[2]
            self._fill_cell_text(abstract_cell, [md.abstract or ""])

        # Keywords → row 2, col 0 (structure: "Keywords:" header, then one kw per para)
        if md.keywords and len(tbl.rows) > 2 and len(tbl.rows[2].cells) > 0:
            kw_cell = tbl.rows[2].cells[0]
            # Accept semicolons OR commas as keyword separator
            kw_list = [k.strip() for k in re.split(r"[;,]", md.keywords) if k.strip()]
            # Preserve "Keywords:" header in para[0], fill paras[1..6], keep DOI block
            cell_paras = kw_cell.paragraphs
            for i, kw in enumerate(kw_list[:6]):
                idx = 1 + i
                if idx < len(cell_paras):
                    for r in list(cell_paras[idx]._element.findall(qn("w:r"))):
                        cell_paras[idx]._element.remove(r)
                    cell_paras[idx].add_run(kw)

    # ----- low-level paragraph helpers -----

    def _new_para(self, *, style: Optional[str] = None, alignment: Optional[str] = None, hanging_in: Optional[float] = None):
        p = self.doc.add_paragraph()
        if style:
            try:
                p.style = style
            except KeyError:
                pass  # style not defined in template, fall through to Normal
        if alignment and alignment in _ALIGN_MAP:
            p.alignment = _ALIGN_MAP[alignment]
        pf = p.paragraph_format
        pf.line_spacing = self.line_spacing
        if hanging_in is not None:
            pf.left_indent = Inches(hanging_in)
            pf.first_line_indent = Inches(-hanging_in)
        return p

    # ----- block renderers -----

    def _render_heading(self, b: HeadingBlock):
        h_spec = self.headings.get(f"Heading {b.level}", {}) or {}
        align = h_spec.get("alignment", "left")
        size = h_spec.get("font_size_pt", self.main_size)
        bold = h_spec.get("bold", b.level == "A")
        italic = h_spec.get("italic", b.level == "C")
        case_rule = h_spec.get("case")

        text = b.text
        if case_rule == "upper":
            text = text.upper()
        elif case_rule == "title_case":
            text = text.title()

        prefix = f"{b.numbering} " if b.numbering else ""
        p = self._new_para(style=f"Heading {b.level}", alignment=align)
        run = p.add_run(prefix + text)
        run.font.name = self.main_font
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    def _render_paragraph(self, b: ParagraphBlock):
        align = b.alignment or "justify"
        p = self._new_para(style=b.style or "Main Text", alignment=align)
        _add_paragraph_text(p, b.text, b.runs, self.main_font, self.main_size)

    def _render_table(self, b: TableBlock):
        if not b.rows and not b.header:
            return
        ncols = len(b.header) if b.header else len(b.rows[0])
        has_header = bool(b.header)
        nrows = (1 if has_header else 0) + len(b.rows)

        table = self.doc.add_table(rows=nrows, cols=ncols)
        table.autofit = True

        borders = self.tables_cfg.get("borders", {}) or {}
        horizontal = borders.get("horizontal", ["top", "below_header", "bottom"]) or []
        want_vertical = borders.get("vertical", False)

        # Header row
        if has_header:
            for i, val in enumerate(b.header):
                cell = table.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.bold = True
                run.font.name = self.table_font
                run.font.size = Pt(self.table_size)
                _set_cell_border(
                    cell,
                    top="top" in horizontal,
                    bottom="below_header" in horizontal,
                    left=want_vertical,
                    right=want_vertical,
                )

        # Body rows
        body_start = 1 if has_header else 0
        for r_idx, row in enumerate(b.rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[body_start + r_idx].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = self.table_font
                run.font.size = Pt(self.table_size)

                is_last_row = (body_start + r_idx) == nrows - 1
                _set_cell_border(
                    cell,
                    top=(r_idx == 0 and not has_header and "top" in horizontal),
                    bottom=(is_last_row and "bottom" in horizontal),
                    left=want_vertical,
                    right=want_vertical,
                )

        if b.caption:
            cap = self._new_para(style="Caption B", alignment="center")
            run = cap.add_run(b.caption)
            run.font.name = self.main_font
            run.font.size = Pt(max(self.table_size, 9))

    def _render_figure(self, b: FigureBlock):
        p = self._new_para(alignment="center")
        if b.src and os.path.exists(b.src):
            run = p.add_run()
            try:
                run.add_picture(b.src, width=Inches(5.5))
            except Exception:
                p.add_run(f"[Figure: {b.alt or b.caption or b.src}]")
        else:
            p.add_run(f"[Figure: {b.alt or b.caption or 'image not found'}]")

        if b.caption:
            cap = self._new_para(style="Caption B", alignment="center")
            run = cap.add_run(b.caption)
            run.font.name = self.main_font
            run.font.size = Pt(max(self.table_size, 9))

    def _render_reference(self, b: ReferenceBlock):
        hanging = self.refs_cfg.get("hanging_indent_in", 0.5)
        p = self._new_para(style="Reference", alignment="left", hanging_in=hanging)
        if b.runs:
            _apply_runs(p, b.runs, self.main_font, self.main_size)
        else:
            run = p.add_run(b.raw)
            run.font.name = self.main_font
            run.font.size = Pt(self.main_size)

    def _render_equation(self, b: EquationBlock):
        p = self._new_para(style="Equation", alignment="right")
        text = b.text or b.latex or ""
        if b.number:
            text = f"{text}    ({b.number})"
        run = p.add_run(text)
        run.italic = True
        run.font.name = self.main_font
        run.font.size = Pt(self.main_size)

    # ----- front matter -----

    def _render_front_matter(self, md):
        tb = self.ruleset.get("title_block", {}) or {}
        title_sz   = (tb.get("title",       {}) or {}).get("font_size_pt", 17)
        title_bold = (tb.get("title",       {}) or {}).get("bold", False)
        author_sz  = (tb.get("author",      {}) or {}).get("font_size_pt", 13)
        affil_sz   = (tb.get("affiliation", {}) or {}).get("font_size_pt",  8)
        abs_sz     = (self.ruleset.get("abstract", {}) or {}).get("font_size_pt", self.main_size)

        if md.title:
            p = self._new_para(style="Title", alignment="center")
            run = p.add_run(md.title)
            run.bold = title_bold
            run.font.name = self.main_font
            run.font.size = Pt(title_sz)

        for author in md.authors:
            p = self._new_para(style="Author", alignment="center")
            run = p.add_run(author)
            run.font.name = self.main_font
            run.font.size = Pt(author_sz)

        for aff in md.affiliations:
            p = self._new_para(style="Affiliation", alignment="center")
            run = p.add_run(aff)
            run.italic = True
            run.font.name = self.main_font
            run.font.size = Pt(affil_sz)

        if md.abstract:
            p = self._new_para(style="Abstract", alignment="justify")
            run = p.add_run("Abstract. ")
            run.bold = True
            run.font.name = self.main_font
            run.font.size = Pt(abs_sz)
            run2 = p.add_run(md.abstract)
            run2.italic = True
            run2.font.name = self.main_font
            run2.font.size = Pt(abs_sz)

        if md.keywords:
            p = self._new_para(alignment="left")
            r1 = p.add_run("Keywords: ")
            r1.italic = True
            r1.font.name = self.main_font
            r1.font.size = Pt(self.main_size)
            r2 = p.add_run(md.keywords)
            r2.font.name = self.main_font
            r2.font.size = Pt(self.main_size)

        if md.doi:
            p = self._new_para(alignment="left")
            r1 = p.add_run("DOI: ")
            r1.italic = True
            r1.font.name = self.main_font
            r1.font.size = Pt(self.main_size)
            r2 = p.add_run(md.doi)
            r2.font.name = self.main_font
            r2.font.size = Pt(self.main_size)

    # ----- public render -----

    def build(self, document) -> Document:
        doc_obj = normalize_input(document)
        if self._has_template:
            self._inject_template_front_matter(doc_obj.metadata)
        else:
            self._render_front_matter(doc_obj.metadata)

        for b in doc_obj.blocks:
            if b.type == "heading":
                self._render_heading(b)
            elif b.type == "paragraph":
                self._render_paragraph(b)
            elif b.type == "table":
                self._render_table(b)
            elif b.type == "figure":
                self._render_figure(b)
            elif b.type == "reference":
                self._render_reference(b)
            elif b.type == "equation":
                self._render_equation(b)

        return self.doc


def render_docx(document: Any, ruleset: Dict[str, Any], output_path: str, template_path: Optional[str] = None) -> str:
    """
    Render a StructuredDocument to a .docx file at output_path.

    If template_path is given (e.g. the UiTM .dotx), its styles, headers,
    and section breaks are used as the base — content is rendered into it.
    """
    renderer = DocxRenderer(ruleset, template_path=template_path)
    doc = renderer.build(document)
    doc.save(output_path)
    return output_path
