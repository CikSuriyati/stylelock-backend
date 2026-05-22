"""
docx_ingest.py — DOCX -> StructuredDocument parser.

The SciSpace-style flow we want to enable:

    user.docx --> ingest_docx() --> StructuredDocument (JSON)
                                            |
                                            v
                            render_docx | render_pdf | render_html
                            (each driven by a ruleset, e.g. mjcet_ruleset.json)

This module is INTENTIONALLY render-agnostic: no CSS, no python-docx style
mutation, no HTML. It only inspects the source document and emits the
structured editorial model.

Style detection strategy
------------------------
1. **Trust the source style first.** If the paragraph already has a style
   name we recognise (Heading A, Heading B, Heading C, Title, Author,
   Affiliation, Abstract, Reference, Caption, FigureCaption, etc.), use it.
   This is the case when authors used the UiTM/MJCET template's Style Gallery.

2. **Heuristic fallback** for paragraphs styled as Normal or with unrecognised
   styles: position in the document, numeric prefix (`1.`, `1.1`, `1.1.1`),
   text content (starts with "Abstract", "Keywords:", "Fig.", "Table",
   "References", contains "@"), and so on. Heuristics here mirror the
   detect_style() logic in main.py's FormattingEngine.

3. **Inline runs** are preserved (bold / italic / underline / superscript /
   subscript) so the editor and renderers can reproduce styled spans.
"""

from __future__ import annotations

import os
import re
from typing import List, Optional, Tuple, Dict, Any

from docx import Document
from docx.document import Document as _DocumentT
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

# Import the canonical schema from renderers/ — same source of truth.
from renderers.schema import (
    StructuredDocument,
    DocumentMetadata,
    HeadingBlock,
    ParagraphBlock,
    TableBlock,
    FigureBlock,
    ReferenceBlock,
    EquationBlock,
    RunSpan,
    Block,
)


# ---------------------------------------------------------------------------
# Style-name normalisation
# ---------------------------------------------------------------------------

# Style names we trust verbatim when they appear in the source DOCX.
# Maps a normalised lowercase name -> canonical name we emit.
KNOWN_STYLE_NAMES = {
    "title": "Title",
    "els-title": "Title",
    "author": "Author",
    "els-author": "Author",
    "affiliation": "Affiliation",
    "els-affiliation": "Affiliation",
    "abstract": "Abstract",
    "els-abstract-text": "Abstract",
    "els-abstract-head": "AbstractHead",
    "heading a": "Heading A",
    "headinga": "Heading A",
    "heading 1": "Heading A",
    "heading b": "Heading B",
    "headingb": "Heading B",
    "heading 2": "Heading B",
    "heading c": "Heading C",
    "headingc": "Heading C",
    "heading 3": "Heading C",
    "main text": "Main Text",
    "maintext": "Main Text",
    "els-body-text": "Main Text",
    "main text [heading a]": "Main Text",
    "maintextheadinga": "Main Text",
    "main text [heading b]": "Main Text",
    "maintextheadingb": "Main Text",
    "main body text [heading c]": "Main Text",
    "mainbodytextheadingc": "Main Text",
    "reference": "Reference",
    "els-reference": "Reference",
    "references": "Reference",
    "caption": "TableCaption",
    "table caption": "TableCaption",
    "els-caption": "TableCaption",
    "figurecaption": "FigureCaption",
    "figure caption": "FigureCaption",
    "footnote": "Footnote",
    "els-footnote": "Footnote",
    "footnotetext": "Footnote",
    "equation": "Equation",
    "els-equation": "Equation",
    # NOTE: "normal" and "default" are intentionally NOT in this map —
    # they're carriers of no information, so we let _infer_style run.
}


def _canonical_style(raw: Optional[str]) -> str:
    """Lowercase + lookup. Returns the canonical name, or '' if unknown
    (including the special-cased 'Normal' / 'Default' which we treat as
    unknown so the heuristic can do its job)."""
    if not raw:
        return ""
    return KNOWN_STYLE_NAMES.get(raw.strip().lower(), "")


# ---------------------------------------------------------------------------
# Heuristic style detection (only used when source style is "Normal"/unknown)
# ---------------------------------------------------------------------------

_RE_HEADING_C = re.compile(r"^\d+\.\d+\.\d+(\.\d+)*[\s\.]")
_RE_HEADING_B = re.compile(r"^\d+\.\d+[\s\.]")
_RE_HEADING_A = re.compile(r"^\d+[\s\.]")
_RE_FIG_CAPTION = re.compile(r"^(fig(?:ure)?\.?\s*\d+\.)", re.IGNORECASE)
_RE_TABLE_CAPTION = re.compile(r"^(table\s*\d+\.)", re.IGNORECASE)


def _infer_style(text: str, index: int, has_title: bool, has_abstract: bool) -> str:
    """Heuristic style for a paragraph whose source style is Normal/unknown.

    `index` is the running count of non-empty paragraphs seen so far (1-based
    after this one is counted). `has_title` and `has_abstract` let the
    heuristic avoid double-claiming front-matter.
    """
    t = text.strip()
    if not t:
        return ""

    low = t.lower()

    # Captions
    if _RE_FIG_CAPTION.match(t):
        return "FigureCaption"
    if _RE_TABLE_CAPTION.match(t):
        return "TableCaption"

    # References block header
    if low in ("references", "reference"):
        return "Heading A"

    # Keywords line — high-confidence prefix
    if low.startswith("keywords:") or low.startswith("keywords "):
        return "Keywords"

    # Abstract header line
    if low == "abstract":
        return "AbstractHead"

    # Numbered headings — most specific first
    if _RE_HEADING_C.match(t):
        return "Heading C"
    if _RE_HEADING_B.match(t):
        return "Heading B"
    if _RE_HEADING_A.match(t):
        return "Heading A"

    # First paragraph is almost always the title (if we haven't claimed one)
    if index == 0 and not has_title:
        return "Title"

    # Front-matter heuristics (within ~6 paragraphs after the title)
    if index < 6 and has_title and not has_abstract:
        # Affiliation indicators are MORE SPECIFIC than author indicators —
        # check them first so "Universiti Teknologi MARA, Shah Alam" wins
        # over the generic comma-author heuristic.
        if any(kw in low for kw in (
            "university", "universiti", "institute", "college",
            "school of", "department", "faculty",
        )):
            return "Affiliation"
        # Author byline indicators: asterisks, commas separating names, "et al"
        if (
            "*" in t
            or " et al" in low
            or (len([c for c in t if c == ","]) >= 1 and len(t.split()) <= 12)
        ):
            return "Author"

    # Abstract heuristic: a substantive paragraph appearing in front matter
    # before any numbered heading has fired. Threshold at 120 chars catches
    # short abstracts without false-positiving short author/affiliation lines.
    if index < 8 and has_title and not has_abstract and len(t) > 120:
        return "Abstract"

    return "Main Text"


# ---------------------------------------------------------------------------
# Run-level extraction (preserve inline formatting)
# ---------------------------------------------------------------------------

def _runs_from_paragraph(p: Paragraph) -> List[RunSpan]:
    """Collapse adjacent runs with identical formatting; drop empty ones."""
    out: List[RunSpan] = []
    for r in p.runs:
        text = r.text or ""
        if not text:
            continue
        vert = None
        try:
            # python-docx exposes vertAlign via run.font.superscript/subscript
            superscript = bool(r.font.superscript)
            subscript = bool(r.font.subscript)
        except Exception:
            superscript = subscript = False

        span = RunSpan(
            text=text,
            bold=bool(r.bold) if r.bold is not None else False,
            italic=bool(r.italic) if r.italic is not None else False,
            underline=bool(r.underline) if r.underline is not None else False,
            superscript=superscript,
            subscript=subscript,
        )

        # Merge with previous if identical formatting (keeps the model small)
        if out and _same_formatting(out[-1], span):
            out[-1] = RunSpan(
                text=out[-1].text + span.text,
                bold=span.bold,
                italic=span.italic,
                underline=span.underline,
                superscript=span.superscript,
                subscript=span.subscript,
            )
        else:
            out.append(span)
    return out


def _same_formatting(a: RunSpan, b: RunSpan) -> bool:
    return (
        a.bold == b.bold
        and a.italic == b.italic
        and a.underline == b.underline
        and a.superscript == b.superscript
        and a.subscript == b.subscript
    )


def _paragraph_alignment(p: Paragraph) -> Optional[str]:
    if p.alignment is None:
        return None
    name = p.alignment.name.lower()
    if name in ("left", "right", "center", "justify"):
        return name
    if name == "justify_low":
        return "justify"
    return None


def _has_image(p: Paragraph) -> bool:
    return bool(p._element.xpath(".//w:drawing"))


# ---------------------------------------------------------------------------
# Table handling
# ---------------------------------------------------------------------------

def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _table_to_block(tbl: Table) -> Tuple[Optional[TableBlock], Optional[str]]:
    """Convert a python-docx Table to a TableBlock.

    Returns (block, abstract_text). If the table is the Article Info /
    Abstract layout table (3 columns, contains an "Abstract" cell), we
    return (None, abstract_text) so the caller routes it into metadata
    instead of emitting a TableBlock.
    """
    cells_text = [[_cell_text(c) for c in row.cells] for row in tbl.rows]

    # Detect the Article Info layout table: 3 columns, has a header row with
    # "ARTICLE INFO" + "ABSTRACT", and an Abstract body in the right column.
    flat_lower = " ".join(c.lower() for row in cells_text for c in row)
    if ("abstract" in flat_lower and "article info" in flat_lower) or (
        any(len(row) == 3 for row in cells_text)
        and "abstract" in flat_lower
    ):
        # Best-effort: take the longest cell in the right-most column as abstract
        abstract_candidates = []
        for row in cells_text:
            if len(row) >= 3 and row[-1]:
                abstract_candidates.append(row[-1])
        if abstract_candidates:
            abstract_text = max(abstract_candidates, key=len)
            return None, abstract_text

    # Ordinary data table
    if not cells_text:
        return None, None
    if len(cells_text) == 1:
        # Single row — treat as header-only table
        return TableBlock(header=cells_text[0], rows=[]), None
    return TableBlock(header=cells_text[0], rows=cells_text[1:]), None


# ---------------------------------------------------------------------------
# Front-matter extraction
# ---------------------------------------------------------------------------

def _extract_keywords(text: str) -> Optional[str]:
    low = text.lower()
    if low.startswith("keywords:"):
        return text.split(":", 1)[1].strip()
    if low.startswith("keywords "):
        return text[len("keywords "):].strip()
    return None


# ---------------------------------------------------------------------------
# Main ingestion
# ---------------------------------------------------------------------------

def ingest_docx(path: str) -> StructuredDocument:
    """Parse a .docx file into a StructuredDocument.

    Raises FileNotFoundError if the path doesn't exist; ValueError if
    python-docx can't open the file.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    try:
        doc: _DocumentT = Document(path)
    except Exception as e:
        raise ValueError(f"Could not open .docx: {e}") from e

    metadata = DocumentMetadata()
    blocks: List[Block] = []

    # Walk the body in document order — paragraphs AND tables, interleaved.
    para_index = 0
    pending_table_caption: Optional[str] = None
    pending_figure_caption: Optional[str] = None
    in_references_section = False  # flips on after a "References" Heading A

    for element in doc.element.body.iterchildren():

        # -------- Paragraph --------
        if isinstance(element, CT_P):
            p = Paragraph(element, doc)
            text = p.text.strip()

            # Inline image? Emit a FigureBlock and consume any pending caption.
            if _has_image(p):
                fig = FigureBlock(caption=pending_figure_caption)
                pending_figure_caption = None
                blocks.append(fig)
                continue

            if not text:
                continue

            # Pick the best style for this paragraph.
            src_style = p.style.name if hasattr(p.style, "name") else None
            canonical = _canonical_style(src_style) or _infer_style(
                text,
                para_index,
                has_title=bool(metadata.title),
                has_abstract=bool(metadata.abstract),
            )
            para_index += 1

            # State machine: once we've crossed into the References section,
            # treat any paragraph that isn't itself a heading as a reference.
            if in_references_section and canonical not in (
                "Heading A", "Heading B", "Heading C",
                "FigureCaption", "TableCaption",
            ):
                canonical = "Reference"

            # ----- Front matter consumed into metadata -----
            if canonical == "Title" and not metadata.title:
                metadata.title = text
                continue

            if canonical == "Author":
                # Author paragraphs may contain multiple authors joined by ",".
                # Keep the raw line; downstream UI can split it.
                metadata.authors.append(text)
                continue

            if canonical == "Affiliation":
                metadata.affiliations.append(text)
                continue

            if canonical == "Abstract":
                metadata.abstract = (
                    (metadata.abstract + " " if metadata.abstract else "") + text
                )
                continue

            if canonical == "Keywords":
                kws = _extract_keywords(text)
                if kws:
                    metadata.keywords = kws
                continue

            if canonical == "AbstractHead":
                # Skip the "Abstract" label paragraph itself
                continue

            # ----- Captions: hold onto them to attach to the next figure/table -----
            if canonical == "FigureCaption":
                pending_figure_caption = text
                continue

            if canonical == "TableCaption":
                pending_table_caption = text
                continue

            # ----- Headings -----
            if canonical in ("Heading A", "Heading B", "Heading C"):
                level = canonical.split()[-1]  # "A" / "B" / "C"
                # Strip leading numbering like "1." or "1.1" — the renderer
                # re-applies it from the ruleset.
                cleaned = re.sub(r"^\d+(\.\d+)*[\s\.]+", "", text).strip() or text
                blocks.append(HeadingBlock(level=level, text=cleaned))
                # Enter references mode after a top-level "References" heading.
                if level == "A" and cleaned.strip().lower().startswith("reference"):
                    in_references_section = True
                # Any other Heading A after that exits references mode
                elif level == "A" and in_references_section:
                    in_references_section = False
                continue

            # ----- References -----
            if canonical == "Reference":
                blocks.append(
                    ReferenceBlock(raw=text, runs=_runs_from_paragraph(p))
                )
                continue

            # ----- Equations -----
            if canonical == "Equation":
                blocks.append(EquationBlock(text=text))
                continue

            # ----- Default: plain paragraph (Main Text, Footnote, anything else) -----
            blocks.append(
                ParagraphBlock(
                    text=text,
                    runs=_runs_from_paragraph(p),
                    style=canonical or "Main Text",
                    alignment=_paragraph_alignment(p),
                )
            )
            continue

        # -------- Table --------
        if isinstance(element, CT_Tbl):
            t = Table(element, doc)
            block, abstract_from_layout = _table_to_block(t)
            if abstract_from_layout and not metadata.abstract:
                metadata.abstract = abstract_from_layout
                continue
            if block is None:
                continue
            block.caption = pending_table_caption
            pending_table_caption = None
            blocks.append(block)
            continue

    return StructuredDocument(metadata=metadata, blocks=blocks)


def ingest_docx_to_dict(path: str) -> Dict[str, Any]:
    """Convenience wrapper: returns a plain dict for direct JSON serialization."""
    return ingest_docx(path).model_dump()
