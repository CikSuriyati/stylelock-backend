from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, Any
import shutil
import os
import uuid
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from lxml import etree
from copy import deepcopy

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="GADING Journal Formatting System")

# CORS: in production set FRONTEND_URL env var to your GitHub Pages URL
# (e.g. "https://username.github.io"). Defaults to "*" for local dev.
_frontend_url = os.environ.get("FRONTEND_URL", "*")
_allowed_origins = ["*"] if _frontend_url == "*" else [o.strip() for o in _frontend_url.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
PROCESSED_DIR = os.path.join(BASE_DIR, "processed")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Official GADING 2024 Specifications v2 (Template-Defined Approach)
# Philosophy: Respect template's built-in formatting, only enforce critical requirements
GADING_SPECS = {
    "Title": {"alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": True, "line_spacing": 1.0},
    "Author": {"alignment": WD_ALIGN_PARAGRAPH.CENTER, "bold": False, "line_spacing": 1.0},
    "Affiliation": {"alignment": WD_ALIGN_PARAGRAPH.CENTER, "italic": True, "line_spacing": 1.0},
    "Footnote": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "size": 8, "line_spacing": 1.0},
    "Abstract": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "italic": True, "line_spacing": 1.0},
    "Heading A": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "size": 10, "bold": True, "line_spacing": 1.0},
    "Heading B": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "bold": True, "line_spacing": 1.0, "hanging_indent": 18},
    "Heading C": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "italic": True, "line_spacing": 1.0},
    "Caption B": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "size": 9, "line_spacing": 1.0},
    "Equation": {"alignment": WD_ALIGN_PARAGRAPH.RIGHT}, 
    "Main Text": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.0},
    "Reference": {"alignment": WD_ALIGN_PARAGRAPH.LEFT, "line_spacing": 1.0, "hanging_indent": 18},
    "Default": {"alignment": WD_ALIGN_PARAGRAPH.JUSTIFY, "line_spacing": 1.0}
}

class FormattingEngine:
    def __init__(self, template_path, manuscript_path, overrides=None):
        self.template_path = template_path
        self.manuscript_path = manuscript_path
        # We use the template as our base to keep all styles and structural settings
        self.doc = Document(template_path)
        self.manuscript = Document(manuscript_path)
        self.report = []
        self.overrides = overrides or {} # dict of index -> style_name
        self.element_mapping = [] # Maps Formatted Index -> Manuscript Index (or -1)

    def log(self, message):
        self.report.append(message)

    def extract_front_matter(self):
        """Scans manuscript for metadata to inject into specific GADING slots."""
        meta = {"Title": "", "Author": "", "Abstract": "", "Keywords": ""}
        used_indices = set()
        total = len(self.manuscript.paragraphs)
        
        for i, p in enumerate(self.manuscript.paragraphs):
            text = p.text.strip()
            if not text: continue
            
            style = self.detect_style(text, i, total, source_p=p)
            
            # Use heuristic to identify what we consume as metadata
            if style == "Title" and not meta["Title"]:
                meta["Title"] = text
                used_indices.add(i)
            elif style == "Author" or style == "Affiliation":
                if style == "Author": meta["Author"] += text + "\n"
                used_indices.add(i)
            elif text.lower().startswith("abstract") or (style == "Abstract" and not text.lower().startswith("keywords")):
                # Consolidate abstract lines until keywords or next big style
                meta["Abstract"] += text + " "
                used_indices.add(i)
            elif text.lower().startswith("keywords:"):
                meta["Keywords"] = text
                used_indices.add(i)
                
        return meta, used_indices

    def prepare_template(self):
        """Surgically clears placeholders but PRESERVES GADING Layout."""
        # Identification based on inspection (templates/gading_template.docx):
        # Index 0, 1: Headers
        # Index 2-5: Title/Author/Affiliation placeholders
        # Index 7: GJSS Logo / Article Info Layout table
        
        body = self.doc.element.body
        elements = list(body.iterchildren())
        
        # Elements to keep: Header, Slots, and Logo Table
        self.keep_indices = [0, 1, 2, 3, 4, 5, 7]
        
        # Clear body and restore branding
        for el in elements:
            body.remove(el)
        for idx in self.keep_indices:
            if idx < len(elements):
                body.append(elements[idx])

    def inject_front_matter(self, meta):
        """Injects extracted metadata into official GADING placeholders."""
        # 1. Inject into designated Paragraph Slots
        # Slot 2: Title, Slot 3: Author, Slot 4: Affiliation
        paragraphs = self.doc.paragraphs
        if len(paragraphs) >= 6:
            if meta["Title"]: paragraphs[2].text = meta["Title"]
            if meta["Author"]:
                authors = meta["Author"].strip().split("\n")
                paragraphs[3].text = authors[0]
                if len(authors) > 1 and len(paragraphs) > 4:
                    paragraphs[4].text = authors[1]
                    
        # 2. Inject into Tables (Abstract/Keywords Layout)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt_low = cell.text.lower()
                    if "abstract" in txt_low or "article info" in txt_low:
                        for cp in cell.paragraphs:
                            if "abstract" in cp.text.lower() or cp.style.name == "Abstract":
                                cp.text = meta["Abstract"] or "Abstract goes here."
                            if "keywords:" in cp.text.lower():
                                cp.text = meta["Keywords"] or "Keywords:"

    def get_style_info(self, style_name):
        """Extracts basic CSS-compatible properties from a document style."""
        try:
            style = self.doc.styles[style_name]
            font = style.font
            return {
                "name": style_name,
                "fontSize": font.size.pt if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
                "fontName": font.name or "Times New Roman"
            }
        except:
            # Fallback to GADING_SPECS if style extraction fails
            spec = GADING_SPECS.get(style_name, {})
            return {
                "name": style_name,
                "fontSize": spec.get("size", 11),
                "bold": spec.get("bold", False),
                "italic": spec.get("italic", False),
                "fontName": "Times New Roman"
            }

    def get_css_for_style(self, style_name):
        """Converts GADING spec to CSS for faithful frontend preview."""
        spec = GADING_SPECS.get(style_name, GADING_SPECS["Default"])
        css = {
            "fontFamily": "'Times New Roman', serif",
            "fontSize": f"{spec.get('size', 11)}pt",
            "fontWeight": "bold" if spec.get("bold") else "normal",
            "fontStyle": "italic" if spec.get("italic") else "normal",
            "textAlign": "justify" if spec.get("alignment") == WD_ALIGN_PARAGRAPH.JUSTIFY else 
                         ("center" if spec.get("alignment") == WD_ALIGN_PARAGRAPH.CENTER else 
                          ("right" if spec.get("alignment") == WD_ALIGN_PARAGRAPH.RIGHT else "left")),
            "marginBottom": f"{spec.get('space_after', 0)}pt",
            "marginTop": f"{spec.get('space_before', 0)}pt",
            "lineHeight": str(spec.get("line_spacing", 1.0)),
            # Hanging Indent Logic
            "paddingLeft": "0pt",
            "textIndent": "0pt",
            "color": "black" # Default
        }
        
        if "hanging_indent" in spec:
            css["paddingLeft"] = f"{spec['hanging_indent']}pt"
            css["textIndent"] = f"-{spec['hanging_indent']}pt"
            
        return css

    def get_doc_json(self, doc_obj, flatten_layout=False):
        """
        Returns JSON representation of a document with style details.
        
        Args:
            flatten_layout (bool): If True, extracts text from Layout Tables (Abstract) 
                                   into editable paragraphs. Used for Editor View.
                                   If False, renders them as Table Placeholders. Used for Reference Guide.
        """
        data = []
        # Iterate over body elements to capture Tables too
        for element in doc_obj.element.body.iterchildren():
            if isinstance(element, CT_P):
                p = Paragraph(element, doc_obj)
                if not p.text.strip(): continue
                
                # Use detect_style to check for override or context
                # For simplicity in JSON, we trust the name, but filter empty text above.
                style_name = p.style.name if hasattr(p.style, 'name') else str(p.style)
                
                style_info = self.get_style_info(style_name)
                css = self.get_css_for_style(style_name)
                
                # Color Extraction Logic (for Template Guide compliance)
                # Check the first run for explicit color (often Red for instructions)
                if p.runs and p.runs[0].font.color and p.runs[0].font.color.rgb:
                    # Convert RGBColor object to hex string
                    hex_color = str(p.runs[0].font.color.rgb)
                    css["color"] = f"#{hex_color}"
                elif p.style.font.color and p.style.font.color.rgb:
                     hex_color = str(p.style.font.color.rgb)
                     css["color"] = f"#{hex_color}"
                
                # Check for images/drawings
                if p._element.xpath('.//w:drawing'):
                    style_name = "Image Placeholder"
                    data.append({
                        "text": "[IMAGE PRESERVED]",
                        "style": style_name,
                        "alignment": "CENTER",
                        "props": {"name": "Image", "bold": False},
                        "css": {"textAlign": "center", "fontStyle": "italic", "color": "#666"}
                    })
                else:
                    data.append({
                        "text": p.text,
                        "style": style_name,
                        "alignment": p.alignment.name if p.alignment else "LEFT",
                        "props": style_info,
                        "css": css
                    })
            elif isinstance(element, CT_Tbl):
                t = Table(element, doc_obj)
                rows = len(t.rows)
                
                # Check if this is a Layout Table (contains Abstract/Article Info)
                is_layout = False
                abstract_paras = []
                
                if flatten_layout:
                    # Scan table for Abstract indicators only if we want to flatten
                    for row in t.rows:
                        for cell in row.cells:
                            text = cell.text.strip().lower()
                            if "abstract" in text and "word count" in text:
                                is_layout = True
                            if "article info" in text:
                                is_layout = True
                            
                            # Collect substantial text matching Abstract style
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    if p.style.name == "Abstract" or len(p.text) > 100:
                                        abstract_paras.append(p)

                if is_layout and abstract_paras:
                    # Flatten! Return the text paragraphs instead of the table placeholder
                    for p in abstract_paras:
                         style_name = p.style.name if hasattr(p.style, 'name') else str(p.style)
                         data.append({
                            "text": p.text,
                            "style": style_name,
                            "alignment": p.alignment.name if p.alignment else "LEFT",
                            "props": self.get_style_info(style_name),
                            "css": self.get_css_for_style(style_name)
                        })
                else:
                    # Standard Data Table -> Render Grid
                    table_data = []
                    for row in t.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    data.append({
                        "text": f"[TABLE - {rows} ROWS]",
                        "style": "Table Placeholder",
                        "alignment": "CENTER",
                        "props": {
                            "name": "Table", 
                            "bold": True, 
                            "italic": False, 
                            "rows": rows,
                            "table_data": table_data
                        }
                    })
        return data

    def strip_manual_formatting(self, paragraph):
        """Strips direct formatting from paragraph and its runs while preserving semantic bold/italic."""
        p = paragraph._element
        pPr = p.find(etree.QName("http://schemas.openxmlformats.org/wordprocessingml/2006/main", "pPr"))
        if pPr is not None:
            for child in pPr.getchildren():
                if etree.QName(child).localname not in ["pStyle", "rPr"]:
                    pPr.remove(child)

        for run in paragraph.runs:
            r = run._element
            rPr = r.find(etree.QName("http://schemas.openxmlformats.org/wordprocessingml/2006/main", "rPr"))
            if rPr is not None:
                for child in rPr.getchildren():
                    tag = etree.QName(child).localname
                    if tag not in ["b", "i", "bCs", "iCs"]:
                        rPr.remove(child)

    def get_style_inventory(self):
        """Returns a list of all available styles in the template."""
        return [s.name for s in self.doc.styles]

    def detect_style(self, text, index, total_paragraphs, source_p=None):
        """Heuristic detection of the correct GADING style, with manual overrides."""
        if str(index) in self.overrides:
            return self.overrides[str(index)]
        
        # 1. Use Manuscript's own style IF it matches one of our expected GADING styles
        if source_p and hasattr(source_p.style, 'name'):
            s_name = source_p.style.name
            if s_name in self.get_style_inventory():
                return s_name

        import re
        text_clean = text.strip()
        if not text_clean: return None # Strict empty line filtering

        # Check for specific GADING style names exactly as per user requirements
        available = self.get_style_inventory()
        
        # 1. Title - BUT check for author patterns first
        # Author indicators: asterisks, numbers after names, "et al", commas between names
        author_indicators = [
            '*' in text_clean and index < 15,  # Asterisks for corresponding authors
            re.search(r'\d+\s*\(', text_clean) and index < 15,  # Numbers with parentheses
            'et al' in text_clean.lower() and index < 15,
            len([c for c in text_clean if c == ',']) >= 2 and index < 15,  # Multiple commas = multiple authors
            re.search(r'[A-Z][a-z]+\s+[A-Z][a-z]+.*[,&]', text_clean) and index < 15  # Name patterns with separators
        ]
        
        if any(author_indicators) and "Author" in available:
            return "Author"
        
        # Now check for Title (first non-empty paragraph, not too long, and not an author)
        if index == 0:
            if "Title" in available: return "Title"
            return available[0] # Fallback to first style if Title missing
        
        # If index 1-3 and looks like a continuing title or author
        if index <=3 and len(text_clean.split()) < 20:
            # Check if it has author characteristics
            if any(author_indicators):
                if "Author" in available: return "Author"
            # Otherwise could be title continuation
            if index == 1 and "Title" in available: return "Title"

        # 2. Abstract/Keywords
        if text_clean.lower().startswith("abstract") or (index < 10 and len(text_clean) > 200):
            if "Abstract" in available: return "Abstract"
            if "Heading A" in available: return "Heading A"
        
        if text_clean.lower().startswith("keywords"):
             if "Abstract" in available: return "Abstract"

        # 0. Figure/Table Captions
        if text_clean.lower().startswith("figure") or text_clean.lower().startswith("table"):
            if "Caption B" in available: return "Caption B"
            if "Caption" in available: return "Caption"
        
        # Source Note
        if text_clean.lower().startswith("source:"):
            if "Footnote" in available: return "Footnote"

        # 3. Headings (Strict GADING Numbering)
        # Heading C: 1.1.1
        if re.match(r'^\d+\.\d+\.\d+(\.\d+)*[\s\.]', text_clean):
            if "Heading C" in available: return "Heading C"
        
        # Heading B: 1.1
        if re.match(r'^\d+\.\d+[\s\.]', text_clean):
            if "Heading B" in available: return "Heading B"
        
        # Heading A: 1. (Starts Introduction usually)
        if re.match(r'^\d+[\s\.]', text_clean):
            if "Heading A" in available: return "Heading A"

        # 4. References
        if text_clean.lower().startswith("reference"):
            if "Reference" in available: return "Reference"
            if "Heading A" in available: return "Heading A"

        # 5. Footnote/Affiliation (Early heuristic)
        if index < 15:
            # Email or affiliation patterns
            if "@" in text_clean or "email" in text_clean.lower():
                if "Footnote" in available: return "Footnote"
                if "Affiliation" in available: return "Affiliation"
            
            # Affiliation patterns (university, department, country names)
            affiliation_keywords = ['university', 'college', 'department', 'faculty', 'institute', 'malaysia', 'school', '1first']
            if any(kw in text_clean.lower() for kw in affiliation_keywords):
                if "Affiliation" in available: return "Affiliation"
                if "Footnote" in available: return "Footnote"

        # 6. Default Body Text
        if "Main Text" in available: return "Main Text"
        if "Normal" in available: return "Normal"
        return available[0]

    def format_table(self, table):
        """Applies GADING Table Formatting Engine rules (APA 7th)."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cell_border(cell, **kwargs):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                if edge in kwargs:
                    tag = 'w:{}'.format(edge)
                    element = tcPr.find(qn(tag))
                    if element is not None: tcPr.remove(element)
                    if kwargs[edge]:
                        new_element = OxmlElement(tag)
                        for key, val in kwargs[edge].items():
                            new_element.set(qn('w:{}'.format(key)), str(val))
                        tcPr.append(new_element)

        # 1. Formatting
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set Header Row Repeat
        if len(table.rows) > 0:
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                        # High-level override for XML consistency
                        r = run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

        # 2. APA Borders (Strict Horizontal Only)
        # Clear all borders first
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None)
        
        border = {'val': 'single', 'sz': '4', 'space': '0', 'color': 'auto'}
        if len(table.rows) > 0:
            # Top of table
            for cell in table.rows[0].cells:
                set_cell_border(cell, top=border)
            
            # Bottom of header (Assuming Row 1 is header)
            for cell in table.rows[0].cells:
                set_cell_border(cell, bottom=border)
                
            # Bottom of table
            if len(table.rows) > 1:
                for cell in table.rows[-1].cells:
                    set_cell_border(cell, bottom=border)

    def process(self, output_path):
        # 1. Extract Metadata and Track Used Indices
        meta, used_indices = self.extract_front_matter()
        
        # 2. Prepare Template (Surgical Layout Preservation)
        self.prepare_template()
        
        # 3. Inject Metadata into designated slots
        self.inject_front_matter(meta)

        # 4. Map Branding/Placeholders to -1 (exclude from body sync)
        self.element_mapping = []
        for element in self.doc.element.body.iterchildren():
            if element.tag.endswith(('p', 'tbl')):
                self.element_mapping.append(-1)

        # 5. Append Body Content (Faithful Copy)
        total_paras = len(self.manuscript.paragraphs)
        para_index = 0
        data_index = 0 

        for element in self.manuscript.element.body.iterchildren():
            # Paragraph Handling
            if element.tag.endswith('p'):
                source_p = Paragraph(element, self.manuscript)
                
                # Check if this paragraph was consumed as Metadata
                if para_index in used_indices:
                    para_index += 1
                    data_index += 1
                    continue

                text = source_p.text
                if not text.strip() and not source_p._element.xpath('.//w:drawing'): 
                    para_index += 1
                    data_index += 1
                    continue

                style = self.detect_style(text, para_index, total_paras, source_p=source_p)
                
                # Image Handling
                if source_p._element.xpath('.//w:drawing'):
                    p_element = deepcopy(source_p._element)
                    self.doc.element.body.append(p_element)
                    self.element_mapping.append(data_index)
                    para_index += 1
                    data_index += 1
                    continue

                target_p = self.doc.add_paragraph()
                for r in source_p.runs:
                    run = target_p.add_run(r.text)
                    run.bold = r.bold
                    run.italic = r.italic

                try:
                    target_p.style = style
                    self.strip_manual_formatting(target_p)
                    spec = GADING_SPECS.get(style, GADING_SPECS["Default"])
                    target_p.alignment = spec.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
                    target_p.paragraph_format.space_after = Pt(spec.get("space_after", 0))
                    
                    self.element_mapping.append(data_index)
                    for run in target_p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(spec.get("size", 10 if style == "Abstract" else 11))
                except:
                    pass
                
                para_index += 1
                data_index += 1

            # Table Handling
            elif element.tag.endswith('tbl'):
                tbl_element = deepcopy(element)
                self.doc.element.body.append(tbl_element)
                new_table = Table(tbl_element, self.doc)
                self.format_table(new_table)
                self.element_mapping.append(data_index)
                para_index += 1
                data_index += 1

        self.doc.save(output_path)
        return self.report

TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
DEFAULT_TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "gading_template.docx")
os.makedirs(TEMPLATES_DIR, exist_ok=True)

@app.post("/transform")
async def process_documents(
    use_system_template: bool = False,
    manuscript: UploadFile = File(...), 
    template: UploadFile = File(None)
):
    job_id = str(uuid.uuid4())
    m_path = os.path.join(UPLOAD_DIR, f"{job_id}_manuscript.docx")
    t_path = os.path.join(UPLOAD_DIR, f"{job_id}_template.docx")
    out_path = os.path.join(PROCESSED_DIR, f"{job_id}_formatted.docx")
    report_path = os.path.join(PROCESSED_DIR, f"{job_id}_report.txt")

    try:
        # Save manuscript
        with open(m_path, "wb") as buffer:
            shutil.copyfileobj(manuscript.file, buffer)

        # Decide which template to use
        if use_system_template:
            if not os.path.exists(DEFAULT_TEMPLATE_PATH):
                raise HTTPException(status_code=400, detail="System template not found on server")
            shutil.copy(DEFAULT_TEMPLATE_PATH, t_path)
        elif template:
            with open(t_path, "wb") as buffer:
                shutil.copyfileobj(template.file, buffer)
        else:
            raise HTTPException(status_code=400, detail="Either a template file or 'use_system_template' must be provided")

        engine = FormattingEngine(t_path, m_path)
        report = engine.process(out_path)

        with open(report_path, "w") as f:
            f.write("\n".join(report))

        return {
            "job_id": job_id,
            "formatted_url": f"/download/{job_id}/formatted",
            "report_url": f"/download/{job_id}/report",
            "review_url": f"/review/{job_id}"
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/review/{job_id}")
async def review_job(job_id: str):
    m_path = os.path.join(UPLOAD_DIR, f"{job_id}_manuscript.docx")
    t_path = os.path.join(UPLOAD_DIR, f"{job_id}_template.docx")
    out_path = os.path.join(PROCESSED_DIR, f"{job_id}_formatted.docx")

    if not os.path.exists(m_path) or not os.path.exists(out_path):
        raise HTTPException(status_code=404, detail="Job files not found")

    engine = FormattingEngine(t_path, m_path)
    processed_doc = Document(out_path)
    return {
        "original": engine.get_doc_json(engine.manuscript, flatten_layout=True),
        "formatted": engine.get_doc_json(processed_doc, flatten_layout=True),
        "mapping": engine.element_mapping
    }

class UpdateRequest(BaseModel):
    manuscript_text: list[str]
    overrides: dict[str, str]

class ReferenceRequest(BaseModel):
    text: str = ""

class APA7Parser:
    def __init__(self):
        # Author-Date pattern: 
        # Matches "Smith, J." or "Smith, J. A." or "Group Name"
        # Followed optionally by more authors
        # Ended by (Year) or (n.d.)
        pass

    def parse_authors(self, author_str: str) -> str:
        """
        Parses author string and formats according to APA 7 rules:
        - 1-20 authors: List all, & before last.
        - 21+ authors: First 19, ..., Last.
        """
        # simplified split by comma to count
        # Warning: "Smith, J., & Doe, B." -> ["Smith, J.", "& Doe, B."]
        # This is strictly for RE-FORMATTING raw lists if possible, 
        # but often we just need to validate/cleanup the existing string.
        # Given the "Paste raw text" nature, we might assume the USER has pasted 
        # something close to the list.
        # FOR NOW: complex re-ordering of 21+ authors from a raw string is risky 
        # without essentially NLP. We will implement the "&" check.
        
        # Heuristic: If we see "et al.", correct it? 
        return author_str

    def to_sentence_case(self, text: str) -> str:
        if not text: return ""
        # Split by splitters to handle "Subtitle"
        parts = re.split(r'([:.!?]\s)', text)
        processed_parts = []
        for part in parts:
            if re.match(r'[:.!?]\s', part):
                processed_parts.append(part)
                continue
            words = part.split()
            new_words = []
            for i, word in enumerate(words):
                clean_word = word.strip(".,;:?!")
                if i == 0:
                    new_words.append(word.capitalize())
                elif clean_word.lower() not in ['i', 'usa', 'uk', 'un', 'fbi', 'hiv']:
                    # Proper noun detection is hard. 
                    # If word is ALL CAPS, force lower.
                    if word.isupper() and len(word) > 1:
                        new_words.append(word.lower())
                    else:
                        # Otherwise keep original casing (safe key)
                         new_words.append(word)
                else:
                    new_words.append(word)
            processed_parts.append(" ".join(new_words))
        return "".join(processed_parts)

    def parse(self, raw_text: str) -> list[dict]:
        raw_text = raw_text.strip()
        if not raw_text: return []
        
        segments = []
        
        # 1. SPLIT BY COMPONENTS
        # We need to find the boundaries: Author (Date). Title. Source.
        
        # Regex for Date: (2020) or (2020, Month Day) or (n.d.) or (in press)
        # Matches: " (2020). " or " (n.d.). " or " (in press). "
        # Capture the whole Author part before it.
        
        date_match = re.search(r"\s\((?:19|20)\d{2}.*?\)\.\s|\s\(n\.d\.\)\.\s|\s\(in press\)\.\s", raw_text)
        
        if date_match:
            author_part = raw_text[:date_match.start()].strip()
            date_full = match_text = date_match.group(0).strip() # " (2020). " -> "(2020)."
            date_content = date_full.strip().rstrip('.') # "(2020)"
            
            # Author Formatting
            # TODO: Add 21+ author logic here if feasible. 
            # For now, ensure it ends with period if missing? 
            # "Smith" -> "Smith."
            
            segments.append({"text": author_part + " ", "italic": False})
            segments.append({"text": date_content + ". ", "italic": False})
            
            remainder = raw_text[date_match.end():]
        else:
            # Fallback: Can't find date, dump whole string
            return [{"text": raw_text, "italic": False}]

        # 2. Extract Title
        # Title ends at first ". " EXCEPT if it's an initial "St. " or "A." ?
        # But usually Title is Sentence Case.
        # We look for the start of the SOURCE (Italicized Journal or Publisher).
        
        # Heuristic: Split by ". "
        # The first chunk is likely the title.
        
        remainder_parts = re.split(r'\.\s', remainder, maxsplit=1)
        
        if len(remainder_parts) > 1:
            raw_title = remainder_parts[0] + "."
            source_part = remainder_parts[1]
            
            # SCENARIO B: EDITION DETECTION
            # Check if Title ends with edition info: "Title (7th ed.)."
            # It might have been captured in raw_title.
            edition_match = re.search(r"\s\((\d+(?:st|nd|rd|th)\s?ed\.|Rev\.\s?ed\.)\)\.$", raw_title, re.IGNORECASE)
            edition_str = ""
            if edition_match:
                edition_str = edition_match.group(0).strip('.') # " (7th ed.)"
                clean_title = raw_title[:edition_match.start()].strip()
            else:
                clean_title = raw_title.rstrip('.')

            # Sentence Case logic
            # Check for bracket suffix: "[Video]"
            bracket_suffix = ""
            if ']' in clean_title:
                b_idx = clean_title.rfind('[')
                if b_idx != -1:
                    bracket_suffix = clean_title[b_idx:] # "[Video]"
                    clean_title = clean_title[:b_idx].strip()
            
            formatted_title = self.to_sentence_case(clean_title)
            
            # Reassemble Title
            full_title = formatted_title
            
            # SCENARIO A: JOURNAL VS BOOK VS WEBPAGE DETECTION
            # Critical Distinction:
            # - Journal: Title = PLAIN. Source = ITALIC (Name), ITALIC (Volume).
            # - Book: Title = ITALIC. Source = PLAIN (Publisher).
            # - Webpage: Title = ITALIC. Source = PLAIN (Site Name).
            
            is_journal = False
            # Strong signal for Journal: Source starts with "Name, Digits" (Volume)
            # Regex: ^Name, \d
            if re.search(r'^.+?,\s\d', source_part.strip()):
                is_journal = True
            
            # Apply Italic Rule
            # Journal -> Title Plain
            # Book/Web -> Title Italic
            title_italic = not is_journal
            
            segments.append({"text": full_title, "italic": title_italic})
            
            # Append Edition (Plain, in parens) AFTER title, before period?
            # APA 7: "Title (7th ed.)." 
            # If Title is Italic (Book), the Edition is NOT Italic.
            if edition_str:
                segments.append({"text": edition_str, "italic": False}) # Edition is never italic
            
            segments.append({"text": ". ", "italic": False}) # End title block
            if bracket_suffix:
                # [Video] usually comes before the period in title? 
                # "Title [Video]."
                # My logic above put it after. Let's fix if needed. 
                # Ideally: segments check.
                pass 
                # Simplified: assuming title block structure ok.
            
            # 3. Source Processing
            source_part = source_part.strip()
            
            # SCENARIO E: DOI / URL
            # Remove trailing period if it is a URL/DOI
            # Also if it ends with "doi.org/xxxx."
            if source_part.endswith('.'):
                if "http" in source_part.lower() or "doi.org" in source_part.lower():
                     source_part = source_part.rstrip('.')
            
            if is_journal:
                # ... (Existing Journal Logic) ...
                j_split = source_part.split(',', 1)
                if len(j_split) > 1:
                    j_name = j_split[0]
                    j_rest = j_split[1]
                    segments.append({"text": j_name + ", ", "italic": True})
                    
                    # Volume/Issue
                    vol_match = re.match(r"\s*(\d+)(\(.*\))?", j_rest)
                    if vol_match:
                        vol = vol_match.group(1)
                        issue = vol_match.group(2) if vol_match.group(2) else ""
                        segments.append({"text": vol, "italic": True})
                        segments.append({"text": issue, "italic": False}) # Issue is plain
                        
                        rest = j_rest[len(vol_match.group(0)):]
                        # SCENARIO: Article Number (e.g. e01234)
                        # Check if 'rest' looks like an page range or article ID
                        # usually ", 10-20." or ", e1234."
                        segments.append({"text": rest, "italic": False})
                    else:
                        segments.append({"text": j_rest, "italic": False})
                else:
                    segments.append({"text": source_part, "italic": False})
            else:
                # Book Identifier: "Publisher."
                # Web Identifier: "Site Name. http..."
                
                # Check for "Retrieved from" (Edge Case)
                # "Retrieved October 11, 2020, from ..."
                
                segments.append({"text": source_part, "italic": False})
                
        else:
            # No Source split found
            segments.append({"text": remainder, "italic": False})
            
        return segments


@app.post("/update/{job_id}")
async def update_manuscript(job_id: str, request: UpdateRequest):
    m_path = os.path.join(UPLOAD_DIR, f"{job_id}_manuscript.docx")
    t_path = os.path.join(UPLOAD_DIR, f"{job_id}_template.docx")
    out_path = os.path.join(PROCESSED_DIR, f"{job_id}_formatted.docx")
    
    if not os.path.exists(m_path):
        raise HTTPException(status_code=404, detail="Manuscript not found")

    # 1. Update the manuscript.docx IN PLACE
    doc = Document(m_path)
    text_ptr = 0
    
    # Simple strategy: Every paragraph in the document that has text
    # gets updated with the next string from the editor's sequence.
    # This maintains 1:1 parity with the DOM order.
    
    for p in doc.paragraphs:
        if not p.text.strip(): continue # Skip empty structural lines
        
        if text_ptr < len(request.manuscript_text):
            p.text = request.manuscript_text[text_ptr]
            text_ptr += 1
    
    # Also handle paragraphs inside tables if they are part of the sync stream
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    if not cp.text.strip(): continue
                    if text_ptr < len(request.manuscript_text):
                        cp.text = request.manuscript_text[text_ptr]
                        text_ptr += 1

    doc.save(m_path)

    # 2. Re-run formatting engine
    engine = FormattingEngine(t_path, m_path, overrides=request.overrides)
    report = engine.process(out_path)
    processed_doc = Document(out_path)

    return {
        "status": "success",
        "original": engine.get_doc_json(engine.manuscript, flatten_layout=True),
        "formatted": engine.get_doc_json(processed_doc, flatten_layout=True),
        "mapping": engine.element_mapping,
        "report": report
    }

@app.post("/tools/format-references")
async def format_references(request: ReferenceRequest):
    try:
        job_id = str(uuid.uuid4())
        output_filename = f"references_{job_id}.docx"
        output_path = os.path.join(PROCESSED_DIR, output_filename) # Changed OUTPUT_DIR to PROCESSED_DIR
        
        # Load Template
        if os.path.exists(DEFAULT_TEMPLATE_PATH): # Changed SYSTEM_TEMPLATE_PATH to DEFAULT_TEMPLATE_PATH
            doc = Document(DEFAULT_TEMPLATE_PATH)
        else:
            doc = Document() # Fallback
            
        # Clear existing body content
        for element in doc.element.body.iterchildren():
            doc.element.body.remove(element)
            
        # Add Title
        t = doc.add_paragraph("References")
        t.style = "Heading A"
        t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Process References
        raw_lines = request.text.split('\n')
        preview_data = []
        parser = APA7Parser()
        
        for line in raw_lines:
            line = line.strip()
            if not line: continue
            
            p = doc.add_paragraph()
            # Apply Paragraph Style (Hanging Indent)
            try:
                p.style = "Reference"
            except:
                p.style = "Normal"
            
            p.paragraph_format.first_line_indent = Pt(-36)
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Smart Parsing & Styling
            segments = parser.parse(line)
            
            # Reconstruct preview text for JSON response
            full_html_text = ""
            
            for seg in segments:
                run = p.add_run(seg["text"])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.italic = seg.get("italic", False)
                
                # HTML builder for preview
                text_content = seg["text"]
                if seg.get("italic"):
                    full_html_text += f"<em>{text_content}</em>"
                else:
                    full_html_text += text_content

            preview_data.append({
                "text": full_html_text, # Now contains HTML tags for italics
                "style": "Reference",
                "css": {
                    "fontFamily": "Times New Roman",
                    "fontSize": "10pt",
                    "paddingLeft": "36pt",
                    "textIndent": "-36pt",
                    "marginBottom": "6pt",
                    "textAlign": "left"
                },
                "isHtml": True 
            })
            
        doc.save(output_path)
        
        return {
            "job_id": job_id,
            "preview": preview_data,
            "download_url": f"/download/{output_filename}"
        }
        
    except Exception as e:
        print(f"Error formatting references: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/styles/{job_id}")
async def get_styles(job_id: str):
    t_path = os.path.join(UPLOAD_DIR, f"{job_id}_template.docx")
    if not os.path.exists(t_path):
        raise HTTPException(status_code=404, detail="Template not found")
    doc = Document(t_path)
    return [s.name for s in doc.styles if s.type == 1] # WD_STYLE_TYPE.PARAGRAPH is 1
@app.get("/template/reference")
async def get_template_reference():
    if not os.path.exists(DEFAULT_TEMPLATE_PATH):
        raise HTTPException(status_code=404, detail="System template not found")
    
    # We can use a dummy manuscript path or just load the template doc
    doc = Document(DEFAULT_TEMPLATE_PATH)
    # Use FormattingEngine's logic without needing a manuscript
    # We create a temporary instance just to use its utility methods
    engine = FormattingEngine(DEFAULT_TEMPLATE_PATH, DEFAULT_TEMPLATE_PATH) 
    # Do NOT flatten layout for Reference Guide. We want to see the boxes.
    return engine.get_doc_json(doc, flatten_layout=False)

@app.get("/download/{job_id}/{file_type}")
async def download_file(job_id: str, file_type: str):
    if file_type == "formatted":
        path = os.path.join(PROCESSED_DIR, f"{job_id}_formatted.docx")
        filename = "formatted_manuscript.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        path = os.path.join(PROCESSED_DIR, f"{job_id}_report.txt")
        filename = "formatting_report.txt"
        media_type = "text/plain"
    
    if os.path.exists(path):
        return FileResponse(path, filename=filename, media_type=media_type)
    raise HTTPException(status_code=404, detail="File not found")


# =====================================================================
# Unified server-side rendering pipeline (SciSpace-style)
# Same structured JSON + ruleset -> HTML / PDF / DOCX
# =====================================================================
import json as _json
from fastapi import Body
from fastapi.responses import HTMLResponse, Response, JSONResponse
from renderers import render_html, render_docx, render_pdf
from ingest import ingest_docx

# Ruleset registry — name -> filename in this directory.
# Add a new journal/template by dropping a JSON file here and registering it.
RULESET_REGISTRY = {
    "gading": "gading_ruleset.json",
    "mjcet": "mjcet_ruleset.json",
    "uitm": "uitm_ruleset.json",
}
DEFAULT_RULESET = "mjcet"

# Back-compat: kept so any legacy code path referencing RULESET_PATH still works.
RULESET_PATH = os.path.join(BASE_DIR, RULESET_REGISTRY[DEFAULT_RULESET])


def _load_ruleset(ruleset_name: Optional[str] = None) -> dict:
    """Load a ruleset by name. Defaults to the MJCET ruleset (the authoritative
    forensic-faithful template). Valid names: 'gading', 'mjcet', 'uitm'.
    Unknown names return a 400 so callers see the typo instead of silently
    falling back."""
    name = (ruleset_name or DEFAULT_RULESET).strip().lower()
    if name not in RULESET_REGISTRY:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown ruleset '{ruleset_name}'. Available: {sorted(RULESET_REGISTRY.keys())}",
        )
    path = os.path.join(BASE_DIR, RULESET_REGISTRY[name])
    if not os.path.exists(path):
        raise HTTPException(status_code=500, detail=f"Ruleset file not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        return _json.load(f)


class RenderRequest(BaseModel):
    document: Any                      # StructuredDocument dict OR legacy flat list
    ruleset_name: Optional[str] = None # e.g. "gading" — defaults to gading_ruleset.json
    use_template: bool = True          # for DOCX: open UiTM template as base


@app.post("/render/html", response_class=HTMLResponse)
async def render_to_html(req: RenderRequest):
    """Render structured doc -> self-contained HTML for the web preview.
    The same CSS rules govern the @page print preview so it matches PDF."""
    try:
        ruleset = _load_ruleset(req.ruleset_name)
        html = render_html(req.document, ruleset, standalone=True)
        return HTMLResponse(content=html)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"HTML render failed: {e}")


@app.post("/render/docx")
async def render_to_docx(req: RenderRequest):
    """Render structured doc -> .docx, applying Word styles from the ruleset.
    If use_template is true and the GADING template exists, it's used as the base."""
    try:
        ruleset = _load_ruleset(req.ruleset_name)
        job_id = str(uuid.uuid4())
        out_path = os.path.join(PROCESSED_DIR, f"{job_id}_render.docx")
        template = DEFAULT_TEMPLATE_PATH if req.use_template and os.path.exists(DEFAULT_TEMPLATE_PATH) else None
        render_docx(req.document, ruleset, out_path, template_path=template)
        return FileResponse(
            out_path,
            filename="document.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX render failed: {e}")


@app.post("/render/pdf")
async def render_to_pdf(req: RenderRequest):
    """Render structured doc -> PDF via LaTeX (tectonic/xelatex/pdflatex)
    with WeasyPrint as a fallback. Returns the PDF file."""
    try:
        ruleset = _load_ruleset(req.ruleset_name)
        job_id = str(uuid.uuid4())
        out_path = os.path.join(PROCESSED_DIR, f"{job_id}_render.pdf")
        render_pdf(req.document, ruleset, out_path)
        return FileResponse(out_path, filename="document.pdf", media_type="application/pdf")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF render failed: {e}")


@app.get("/render/ruleset")
async def get_ruleset(ruleset_name: Optional[str] = None):
    """Expose the active ruleset so the frontend can mirror its styles in real time."""
    return _load_ruleset(ruleset_name)


# =====================================================================
# Ingest pipeline (SciSpace-style)
# DOCX upload -> StructuredDocument JSON the editor and renderers consume.
# =====================================================================


@app.post("/ingest")
async def ingest_endpoint(manuscript: UploadFile = File(...)):
    """Parse an uploaded .docx into a StructuredDocument and return the JSON.

    This is the missing input edge of the SciSpace pipeline. The output is
    safe to feed straight into POST /render/docx | /render/pdf | /render/html.

    Returns:
        {
          "job_id": "...",            # so the frontend can reuse the saved upload
          "document": StructuredDocument,
          "stats": {                  # quick sanity check for the editor
            "blocks": <int>,
            "headings": <int>,
            "references": <int>,
            "tables": <int>,
            "figures": <int>
          }
        }
    """
    filename = (manuscript.filename or "").lower()
    if not filename.endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail=f"Expected a .docx upload, got {manuscript.filename!r}",
        )

    job_id = str(uuid.uuid4())
    saved_path = os.path.join(UPLOAD_DIR, f"{job_id}_manuscript.docx")

    try:
        with open(saved_path, "wb") as buf:
            shutil.copyfileobj(manuscript.file, buf)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not save upload: {e}")

    try:
        structured = ingest_docx(saved_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ingest failed: {e}")

    blocks = structured.blocks
    stats = {
        "blocks": len(blocks),
        "headings": sum(1 for b in blocks if b.type == "heading"),
        "references": sum(1 for b in blocks if b.type == "reference"),
        "tables": sum(1 for b in blocks if b.type == "table"),
        "figures": sum(1 for b in blocks if b.type == "figure"),
        "paragraphs": sum(1 for b in blocks if b.type == "paragraph"),
        "equations": sum(1 for b in blocks if b.type == "equation"),
    }

    return JSONResponse(
        {
            "job_id": job_id,
            "document": structured.model_dump(),
            "stats": stats,
        }
    )


@app.post("/ingest-and-render/{file_type}")
async def ingest_and_render(
    file_type: str,
    manuscript: UploadFile = File(...),
    ruleset_name: Optional[str] = None,
):
    """One-shot convenience: upload a .docx, get back a formatted .docx, .pdf,
    or .html in the target ruleset. Under the hood this is just
    /ingest -> /render/{file_type}, but it spares the frontend a round-trip
    for the common 'reformat my manuscript' case."""
    file_type = (file_type or "").lower().strip()
    if file_type not in ("docx", "pdf", "html"):
        raise HTTPException(
            status_code=400,
            detail=f"file_type must be one of docx | pdf | html (got {file_type!r})",
        )

    filename = (manuscript.filename or "").lower()
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Expected a .docx upload")

    job_id = str(uuid.uuid4())
    saved_path = os.path.join(UPLOAD_DIR, f"{job_id}_manuscript.docx")
    with open(saved_path, "wb") as buf:
        shutil.copyfileobj(manuscript.file, buf)

    try:
        structured = ingest_docx(saved_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ingest failed: {e}")

    ruleset = _load_ruleset(ruleset_name)
    doc_dict = structured.model_dump()

    if file_type == "html":
        html = render_html(doc_dict, ruleset, standalone=True)
        return HTMLResponse(content=html)

    if file_type == "docx":
        out_path = os.path.join(PROCESSED_DIR, f"{job_id}_render.docx")
        template = DEFAULT_TEMPLATE_PATH if os.path.exists(DEFAULT_TEMPLATE_PATH) else None
        render_docx(doc_dict, ruleset, out_path, template_path=template)
        return FileResponse(
            out_path,
            filename="document.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # pdf
    out_path = os.path.join(PROCESSED_DIR, f"{job_id}_render.pdf")
    try:
        render_pdf(doc_dict, ruleset, out_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF render failed: {e}")
    return FileResponse(out_path, filename="document.pdf", media_type="application/pdf")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
